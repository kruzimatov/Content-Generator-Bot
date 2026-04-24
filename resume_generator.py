import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ConversationHandler

from utils import ai_request, send_long_message

# Conversation states
CV_NAME, CV_JOB, CV_EXPERIENCE, CV_SKILLS, CV_PHOTO, CV_TEMPLATE, CV_LANGUAGE = range(20, 27)

# Store user selections per chat
user_data = {}

# Template configs: (font_name, header_color, title_align)
TEMPLATE_CONFIGS = {
    "modern": ("Helvetica", RGBColor(0, 102, 204), WD_ALIGN_PARAGRAPH.LEFT),
    "classic": ("Times New Roman", RGBColor(0, 0, 0), WD_ALIGN_PARAGRAPH.CENTER),
    "creative": ("Georgia", RGBColor(153, 0, 0), WD_ALIGN_PARAGRAPH.RIGHT),
}


def generate_resume_docx(text: str, name: str, template: str, photo_path: str = None) -> str:
    doc = Document()
    _, main_color, title_align = TEMPLATE_CONFIGS.get(
        template, TEMPLATE_CONFIGS["modern"]
    )

    p = doc.add_paragraph()
    p.alignment = title_align
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = main_color

    if photo_path and os.path.exists(photo_path):
        try:
            doc.add_picture(photo_path, width=Inches(1.3))
        except Exception:
            pass

    if template == "classic":
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep.add_run("_" * 60)

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped and stripped.endswith(":") and len(stripped) < 60:
            p = doc.add_paragraph()
            p.space_before = Pt(16)
            p.space_after = Pt(4)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = main_color
            if template == "classic":
                sep = doc.add_paragraph()
                sep.add_run("-" * 40)
                sep.space_after = Pt(2)
        elif stripped:
            doc.add_paragraph(stripped)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", prefix=f"Resume_{name}_")
    doc.save(tmp.name)
    return tmp.name


# --- Handlers ---

async def cv_name(update: Update, context) -> int:
    chat_id = update.message.chat_id
    user_data[chat_id] = {"cv_name": update.message.text.strip()}
    await update.message.reply_text(
        "✅ Ism qabul qilindi!\n\n"
        "💼 Kasbingizni yozing:\n"
        "Masalan: Software Engineer, Dizayner, Marketing Manager..."
    )
    return CV_JOB


async def cv_job(update: Update, context) -> int:
    chat_id = update.message.chat_id
    user_data[chat_id]["cv_job"] = update.message.text.strip()
    await update.message.reply_text(
        "✅ Kasb qabul qilindi!\n\n"
        "📋 Tajribangizni yozing:\n"
        "Masalan: 3 yil Python developer, 2 yil freelance dizayner..."
    )
    return CV_EXPERIENCE


async def cv_experience(update: Update, context) -> int:
    chat_id = update.message.chat_id
    user_data[chat_id]["cv_experience"] = update.message.text.strip()
    await update.message.reply_text(
        "✅ Tajriba qabul qilindi!\n\n"
        "🛠 Skillaringizni yozing (vergul bilan ajrating):\n"
        "Masalan: Python, JavaScript, Figma, Communication, Leadership..."
    )
    return CV_SKILLS


async def cv_skills(update: Update, context) -> int:
    chat_id = update.message.chat_id
    user_data[chat_id]["cv_skills"] = update.message.text.strip()

    keyboard = [
        [InlineKeyboardButton("📸 Rasm yuborish", callback_data="cv_photo_yes")],
        [InlineKeyboardButton("⏭ O'tkazib yuborish", callback_data="cv_photo_skip")],
    ]
    await update.message.reply_text(
        "✅ Skillar qabul qilindi!\n\n"
        "📷 Resume uchun rasmingizni yubormoqchimisiz?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return CV_PHOTO


async def cv_photo_choice(update: Update, context) -> int:
    query = update.callback_query
    await query.answer()
    chat_id = query.message.chat_id

    if query.data == "cv_photo_yes":
        await query.edit_message_text(
            "📷 Rasmingizni yuboring:\n"
            "(Oddiy rasm sifatida yuboring, file emas)"
        )
        return CV_PHOTO
    else:
        user_data[chat_id]["cv_photo"] = None
        return await _ask_template(query.message, chat_id, edit=True)


async def cv_photo_receive(update: Update, context) -> int:
    chat_id = update.message.chat_id
    photo = update.message.photo[-1]
    file = await photo.get_file()

    photo_path = os.path.join(tempfile.gettempdir(), f"cv_photo_{chat_id}.jpg")
    await file.download_to_drive(photo_path)
    user_data[chat_id]["cv_photo"] = photo_path

    await update.message.reply_text("✅ Rasm qabul qilindi!")
    return await _ask_template(update.message, chat_id, edit=False)


async def _ask_template(message, chat_id, edit=False):
    keyboard = [
        [
            InlineKeyboardButton("🔵 Modern", callback_data="cv_tpl_modern"),
            InlineKeyboardButton("⚫ Classic", callback_data="cv_tpl_classic"),
        ],
        [
            InlineKeyboardButton("🔴 Creative", callback_data="cv_tpl_creative"),
        ],
    ]
    text = (
        "🎨 Resume shablonini tanlang:\n\n"
        "🔵 Modern — zamonaviy, ko'k rangda\n"
        "⚫ Classic — an'anaviy, qora-oq\n"
        "🔴 Creative — ijodiy, qizil rangda"
    )
    if edit:
        await message.edit_text(text, reply_markup=InlineKeyboardMarkup(keyboard))
    else:
        await message.reply_text(text, reply_markup=InlineKeyboardMarkup(keyboard))
    return CV_TEMPLATE


async def cv_template(update: Update, context) -> int:
    query = update.callback_query
    await query.answer()
    chat_id = query.message.chat_id

    tpl_map = {
        "cv_tpl_modern": "modern",
        "cv_tpl_classic": "classic",
        "cv_tpl_creative": "creative",
    }
    user_data[chat_id]["cv_template"] = tpl_map.get(query.data, "modern")

    keyboard = [
        [
            InlineKeyboardButton("🇺🇿 O'zbekcha", callback_data="cv_lang_uz"),
            InlineKeyboardButton("🇬🇧 English", callback_data="cv_lang_en"),
        ],
    ]
    await query.edit_message_text(
        f"🎨 Shablon: {user_data[chat_id]['cv_template'].title()}\n\n"
        "🌐 Resume qaysi tilda bo'lsin?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return CV_LANGUAGE


async def cv_language(update: Update, context) -> int:
    query = update.callback_query
    await query.answer()
    chat_id = query.message.chat_id

    lang = "O'zbek" if query.data == "cv_lang_uz" else "English"
    user_data[chat_id]["cv_language"] = lang
    data = user_data[chat_id]
    template = data.get("cv_template", "modern")
    photo_path = data.get("cv_photo")

    await query.edit_message_text("⏳ Resume generatsiya qilinmoqda...")

    prompt = (
        f"You are a professional resume writer.\n\n"
        f"Create a structured and professional resume using:\n\n"
        f"Name: {data['cv_name']}\n"
        f"Job: {data['cv_job']}\n"
        f"Experience: {data['cv_experience']}\n"
        f"Skills: {data['cv_skills']}\n"
        f"Language: {lang}\n\n"
        f"Include:\n"
        f"1. Professional summary (3-4 sentences)\n"
        f"2. Skills section (organized by category)\n"
        f"3. Work experience (with bullet points)\n"
        f"4. Education section\n\n"
        f"Also generate:\n"
        f"- A cover letter (3 paragraphs)\n"
        f"- A short resume version (1 page summary)\n\n"
        f"FORMATTING RULES:\n"
        f"- Do NOT use markdown symbols: **, ##, ###, ---, ``` are FORBIDDEN\n"
        f"- Use only plain text and emoji\n"
        f"- Separate sections with emoji headers like:\n\n"
        f"PROFILE SUMMARY:\n"
        f"(summary text)\n\n"
        f"SKILLS:\n"
        f"(skills list)\n\n"
        f"WORK EXPERIENCE:\n"
        f"(experience details)\n\n"
        f"EDUCATION:\n"
        f"(education details)\n\n"
        f"COVER LETTER:\n"
        f"(cover letter text)\n\n"
        f"SHORT RESUME:\n"
        f"(short version)\n\n"
        f"Keep it clear, modern, and well-structured.\n"
        f"Write everything in {lang} language."
    )

    try:
        result = ai_request(
            f"You are a professional resume writer and career consultant. Write everything in {lang}.",
            prompt,
        )
    except Exception as e:
        result = f"❌ Xatolik yuz berdi: {e}"

    await send_long_message(query.message, result)

    name_safe = data["cv_name"].replace(" ", "_")
    docx_path = None

    try:
        docx_path = generate_resume_docx(result, data["cv_name"], template, photo_path)
        with open(docx_path, "rb") as f:
            await query.message.reply_document(
                document=f,
                filename=f"Resume_{name_safe}.docx",
                caption="📝 Resume — Word (DOCX) formatda",
            )
    except Exception as e:
        await query.message.reply_text(f"⚠️ DOCX yaratishda xatolik: {e}")

    if docx_path:
        try:
            os.unlink(docx_path)
        except OSError:
            pass
    if photo_path:
        try:
            os.unlink(photo_path)
        except OSError:
            pass

    await query.message.reply_text("🔄 Qaytadan boshlash uchun /start bosing!")
    user_data.pop(chat_id, None)
    return ConversationHandler.END
